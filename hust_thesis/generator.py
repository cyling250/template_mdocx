"""
HUST 博士论文生成器 — 基于 template.docx
=========================================
  仅含模板样式映射 + 编排逻辑，所有格式化委托 common 包。
"""
import os, re
from docx import Document

from common import (
    parse_md, parse_bib,
    BibFormatter, CitationFormatter,
    apply_font, render_bibliography,
    add_ref_field,
    make_three_line_table, make_formula_block, make_image_block,
    RE_REF, RE_CITE, RE_INLINE,
)
from .formatters import HustBibFormatter, HustCitationFormatter

_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template.docx")

# ═══ 模板样式映射 ═══
_STYLE_MAP = {
    "body": "Normal", "heading_1": "Heading 1", "heading_2": "Heading 2",
    "heading_3": "Heading 3", "heading_4": "Heading 4",
    "caption": "Caption", "bibliography": "Bibliography",
}
_TABLE_CELL_STYLE = "af3"   # 表格
_IMAGE_STYLE = "af4"        # 图片
_STYLEREF_LEVEL = 1


# ═══════════════════════════════════════════
#  主函数
# ═══════════════════════════════════════════

def generate_thesis(md_content, output_path, template_path=None, bib_path=None,
                    bib_formatter: BibFormatter | None = None,
                    citation_formatter: CitationFormatter | None = None):
    if template_path is None:
        template_path = _TEMPLATE_PATH
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"模板不存在: {template_path}")
    if bib_formatter is None:
        bib_formatter = HustBibFormatter()
    if citation_formatter is None:
        citation_formatter = HustCitationFormatter()

    blocks = parse_md(md_content)

    # ═══ 严格校验 ═══
    errors = []
    ch = 0; tb = 0; eq = 0
    refs = {}
    for b in blocks:
        if b["type"] == "heading" and b["level"] == 1:
            ch += 1; eq = 0; tb = 0
        elif b["type"] == "figure":
            refs[b["label"]] = f"fig_{b['label']}"
        elif b["type"] == "table_with_name":
            tb += 1; refs[b["label"]] = f"tbl_{ch}_{tb}"
        elif b["type"] == "table":
            errors.append(f"第{ch}章第{tb+1}个表格缺少表名，请使用 @table[标签]{{表名}} 指定")
            tb += 1
        elif b["type"] == "formula_with_label":
            eq += 1; refs[b["label"]] = f"eq_{ch}_{eq}"
        elif b["type"] == "formula_missing_label":
            errors.append(f"第{ch}章第{eq+1}个块公式缺少标签，请使用 @formula[标签] 声明")

    all_refs = set()
    for b in blocks:
        if b["type"] == "para":
            for m in RE_REF.finditer(b["text"]):
                all_refs.add(m.group(1).replace('\\_', '_'))
    for label in all_refs:
        if label not in refs:
            errors.append(f"交叉引用 @ref{{{label}}} 找不到对应的图/表标签")
    if errors:
        raise ValueError("生成失败，请修复以下问题后重试：\n  " + "\n  ".join(errors))

    # ── 预扫描 BIB → 构建引用映射 ──
    import shutil
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    cite_map = {}
    for b in blocks:
        if b["type"] == "bib":
            bib_file = _resolve_bib(b.get("path_hint"), bib_path, template_path)
            entries = parse_bib(bib_file) if bib_file else []
            for idx, entry in enumerate(entries, 1):
                cite_map[entry["key"]] = f"bib_{idx}"
            break

    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    # ── 渲染 ──
    ch = 0; tb = 0; eq = 0
    for b in blocks:
        t = b["type"]

        if t == "heading":
            lv = b["level"]; txt = b["text"]
            if lv == 1: ch += 1; eq = 0; tb = 0
            p = doc.add_paragraph()
            try: p.style = doc.styles[_STYLE_MAP.get(f"heading_{lv}", _STYLE_MAP["body"])]
            except: pass
            p.add_run(txt)

        elif t == "para":
            txt = _resolve_refs(b["text"], refs, cite_map, citation_formatter)
            p = doc.add_paragraph()
            try: p.style = doc.styles[_STYLE_MAP["body"]]
            except: pass
            for seg in re.split(r'(@@RF@@[^@]+@@|@@CTX@@[^@]+@@|@@CT@@[^@]+@@)', txt):
                if not seg: continue
                if seg.startswith('@@RF@@') and seg.endswith('@@'):
                    add_ref_field(p, seg[6:-2])
                elif seg.startswith('@@CT@@') and seg.endswith('@@'):
                    add_ref_field(p, seg[6:-2],
                                  number_only=True,
                                  superscript=citation_formatter.ref_superscript)
                elif seg.startswith('@@CTX@@') and seg.endswith('@@'):
                    txt = seg[7:-2]
                    r = p.add_run(txt)
                    r.font.superscript = True
                else:
                    for fp in RE_INLINE.split(seg):
                        if not fp: continue
                        if fp.startswith("**") and fp.endswith("**"):
                            r = p.add_run(fp[2:-2]); apply_font(r, bold=True, cn="黑体")
                        elif fp.startswith("*") and fp.endswith("*"):
                            r = p.add_run(fp[1:-1]); apply_font(r, italic=True, cn="楷体")
                        elif fp.startswith("$") and fp.endswith("$"):
                            from common.formulas import add_omath_inline
                            add_omath_inline(p, fp[1:-1])
                        else: p.add_run(fp)

        elif t == "figure":
            src = b.get("src")
            bm = refs.get(b["label"], f"fig_{b['label']}")
            make_image_block(doc, src, image_style_id=_IMAGE_STYLE,
                             bookmark_name=bm, caption_text=b["caption"],
                             caption_style_name=_STYLE_MAP["caption"],
                             styleref_level=_STYLEREF_LEVEL)

        elif t == "table_with_name":
            tb += 1
            bm = f"tbl_{ch}_{tb}"
            make_three_line_table(doc, b["rows"], _TABLE_CELL_STYLE,
                                  bookmark_name=bm, caption_text=b["desc"],
                                  caption_style_name=_STYLE_MAP["caption"],
                                  styleref_level=_STYLEREF_LEVEL)

        elif t == "table":
            tb += 1
            bm = f"tbl_{ch}_{tb}"
            make_three_line_table(doc, b["rows"], _TABLE_CELL_STYLE,
                                  bookmark_name=bm,
                                  caption_style_name=_STYLE_MAP["caption"],
                                  styleref_level=_STYLEREF_LEVEL)

        elif t == "formula_with_label":
            eq += 1
            make_formula_block(doc, b["text"], ch, eq,
                               bookmark_name=f"eq_{ch}_{eq}",
                               styleref_level=_STYLEREF_LEVEL)

        elif t == "bib":
            bib_file = _resolve_bib(b.get("path_hint"), bib_path, template_path)
            entries = parse_bib(bib_file) if bib_file else []
            cite_map = render_bibliography(
                doc, entries,
                bibliography_style_name=_STYLE_MAP["bibliography"],
                formatter=bib_formatter)

    doc.save(output_path)
    return output_path


# ── 辅助 ──

def _resolve_refs(text, refs, cite_map, citation_formatter):
    def _rep_ref(m):
        label = m.group(1).replace('\\_', '_')
        bm = refs.get(label)
        return f"@@RF@@{bm}@@" if bm else f"[{label}]"

    def _rep_cite(m):
        keys = [k.strip() for k in m.group(1).split(",")]
        return citation_formatter.format_keys(keys, cite_map)

    text = RE_REF.sub(_rep_ref, text)
    text = RE_CITE.sub(_rep_cite, text)
    return text


def _resolve_bib(path_hint, bib_path, template_path):
    # 优先使用外部显式传入的 bib_path（如果存在）
    if bib_path and os.path.isfile(bib_path):
        return bib_path
    # 否则按 path_hint 搜索
    bib_file = path_hint or bib_path
    if not bib_file: return None
    if os.path.isfile(bib_file): return bib_file
    for base in [os.getcwd(),
                 os.path.dirname(template_path),
                 os.path.dirname(os.path.dirname(template_path))]:
        alt = os.path.join(base, bib_file)
        if os.path.isfile(alt): return alt
        alt_docs = os.path.join(base, "docs", bib_file)
        if os.path.isfile(alt_docs): return alt_docs
    return None
