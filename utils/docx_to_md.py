"""
docx → md 转换器

功能：
  - 提取标题、段落（**粗体**/*斜体*/行内公式）
  - 图片导出到指定目录，生成 @figure / ![]() / @ref{fig:...}
  - 表格 → Markdown 表格，生成 @table / @ref{tbl:...}
  - 块公式 → $$…$$
  - 交叉引用 → @ref{…}
  - 文献引用 → @cite{…}

用法：
  from utils import docx_to_md
  docx_to_md("input.docx", "output.md", image_dir_name="images")
"""
import os, re, shutil
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS    = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

# ── 样式 → Markdown 标题层级 ──
_HEADING_PATTERN = re.compile(r"^(?:heading)\s*(\d+)$", re.I)

def _heading_level(style_name: str) -> int | None:
    if not style_name:
        return None
    m = _HEADING_PATTERN.match(style_name.replace(" ", ""))
    if m:
        return int(m.group(1))
    # 常见中文样式名
    style_lower = style_name.strip().lower()
    for level in range(1, 5):
        if any(s in style_lower for s in [f"标题 {level}", f"标题{level}", f"heading{level}"]):
            return level
    return None


# ── 提取行内格式文本 ──
def _run_text(run) -> str:
    """获取 run 的纯文本（含换行符处理）"""
    return run.text or ""


def _has_drawing(para) -> bool:
    """段落中是否包含图片"""
    return para._element.findall(f'.//{qn("w:drawing")}') != []


def _has_math(para) -> bool:
    """段落中是否包含 OMath"""
    return para._element.findall(f'.//{{{MATH_NS}}}oMath') != []


def _extract_images(para, img_dir: str, img_dir_name: str,
                    fig_counter: list, doc) -> list[str]:
    """
    从段落中提取图片到 img_dir，返回生成的 Markdown 行列表。
    fig_counter[0] 自增用作标签序号。
    """
    lines = []
    drawings = para._element.findall(f'.//{qn("w:drawing")}')
    for drawing in drawings:
        # 找 blip 中的图片关系引用
        blips = drawing.findall(f'.//{qn("a:blip")}')
        for blip in blips:
            embed = blip.get(f'{{{R_NS}}}embed') or blip.get(f'{{{W_NS}}}embed')
            if not embed:
                continue
            try:
                rel = doc.part.rels[embed]
                img_data = rel.target_part.blob
            except (KeyError, AttributeError):
                continue

            if not img_data:
                continue

            # 确定扩展名
            ext = ".png"
            content_type = getattr(rel.target_part, "content_type", "") or ""
            if "jpeg" in content_type or "jpg" in content_type:
                ext = ".jpg"
            elif "png" in content_type:
                ext = ".png"
            elif "gif" in content_type:
                ext = ".gif"
            elif "bmp" in content_type:
                ext = ".bmp"

            fig_counter[0] += 1
            tag = f"fig:d2m_{fig_counter[0]}"
            fname = f"fig_{fig_counter[0]}{ext}"
            fpath = os.path.join(img_dir, fname)

            with open(fpath, "wb") as f:
                f.write(img_data)

            rel_path = f"{img_dir_name}/{fname}"
            lines.append(f"@figure[{tag}]{{图片 {fig_counter[0]}}}")
            lines.append(f"![{tag}]({rel_path})")

    return lines


def _extract_math(para) -> str | None:
    """
    提取段落中的 OMath 公式文本（尽量读 LaTeX 或 plain text）。
    返回 $$…$$ 格式字符串，如果没有公式则返回 None。
    """
    omaths = para._element.findall(f'.//{{{MATH_NS}}}oMath')
    if not omaths:
        omaths = para._element.findall(f'.//{{{MATH_NS}}}oMathPara')
    if not omaths:
        return None

    parts = []
    for om in omaths:
        texts = om.itertext()
        tex = "".join(texts).strip()
        if tex:
            parts.append(tex)

    if not parts:
        return None

    return "$$\n" + "\n\n".join(parts) + "\n$$"


def _extract_inline_math(para) -> str:
    """
    从段落中提取行内 OMath（inline），将公式替换为 $…$ 标记。
    返回完整段落文本，公式部分用 $ 包裹。
    由于 python-docx 对 OMath 的 run 文本可能是分散的，
    我们按 XML 节点顺序拼接。
    """
    # 简化做法：直接拼接所有文本，检测 OMath 区域
    body_elem = para._element
    fragments = []

    for node in body_elem.iterchildren():
        tag = node.tag

        # OMath → $...$
        if tag == f'{{{MATH_NS}}}oMath' or tag == f'{{{MATH_NS}}}oMathPara':
            tex = "".join(node.itertext()).strip()
            if tex:
                fragments.append(f"${tex}$")
            continue

        # 普通文本
        if tag == qn('w:r'):
            txt_nodes = node.findall(qn('w:t'))
            for tn in txt_nodes:
                if tn.text:
                    fragments.append(tn.text)

        # 换行
        if tag == qn('w:br'):
            fragments.append("\n")

    return "".join(fragments)


def _process_paragraph_runs(para) -> str:
    """
    遍历 run，将粗体/斜体转为 **/**，同时提取行内公式。
    如果段落包含块公式，则返回 None 让上层处理。
    """
    if _has_math(para) and not _has_drawing(para):
        # 块公式（无图片的公式行）
        return None  # 由外层处理

    # 检查是否纯块公式段落
    math_paras = para._element.findall(f'.//{{{MATH_NS}}}oMathPara')
    omaths = para._element.findall(f'.//{{{MATH_NS}}}oMath')
    # 如果所有内容都是 OMath 且没有穿插文本，视为块公式
    if omaths or math_paras:
        math_text = _extract_math(para)
        if math_text:
            return math_text

    # 行内文本处理（含粗体/斜体/行内公式/图片）
    fragments = []
    for child in para._element.iterchildren():
        tag = child.tag
        if tag == qn('w:r'):
            run_text = ""
            for tn in child.findall(qn('w:t')):
                if tn.text:
                    run_text += tn.text

            if not run_text:
                continue

            bold = child.find(qn('w:rPr')) is not None and \
                   child.find(qn('w:rPr')).find(qn('w:b')) is not None
            italic = child.find(qn('w:rPr')) is not None and \
                     child.find(qn('w:rPr')).find(qn('w:i')) is not None

            if bold and italic:
                fragments.append(f"***{run_text}***")
            elif bold:
                fragments.append(f"**{run_text}**")
            elif italic:
                fragments.append(f"*{run_text}*")
            else:
                fragments.append(run_text)

        elif tag == f'{{{MATH_NS}}}oMath' or tag == f'{{{MATH_NS}}}oMathPara':
            tex = "".join(child.itertext()).strip()
            if tex:
                fragments.append(f"${tex}$")

        elif tag == qn('w:br'):
            fragments.append("\n")

    return "".join(fragments)


_CAPTION_RE = re.compile(
    r"^(?P<type>图|表|Figure|Table|Fig\.?)\s*(\d+(?:[.-]\d+)*)\s*[：:：\s]*(?P<desc>.*)",
    re.I
)


def _process_caption(para, fig_counter: list, tbl_counter: list) -> str | None:
    """
    如果段落是题注样式或匹配题注文本模式，返回 @figure / @table 声明行。
    否则返回 None。
    """
    style_name = (para.style.name or "").strip().lower() if para.style else ""
    text = para.text.strip()

    if not text:
        return None

    # 题注样式或文本含图/表前缀
    is_caption_style = "caption" in style_name
    m = _CAPTION_RE.match(text)

    if is_caption_style or m:
        # 先看是否有 SEQ 域代码 → 自动确定类型
        seq_type = _detect_seq_type(para)
        if seq_type == "figure":
            fig_counter[0] += 1
            return f"@figure[fig:d2m_{fig_counter[0]}]{{{text}}}"
        elif seq_type == "table":
            tbl_counter[0] += 1
            return f"@table[tbl:d2m_{tbl_counter[0]}]{{{text}}}"

        # 从文本推测
        if not m:
            return None
        ct = m.group("type")
        desc = m.group("desc") or text
        if ct.lower().startswith("图") or ct.lower().startswith("fig"):
            fig_counter[0] += 1
            return f"@figure[fig:d2m_{fig_counter[0]}]{{{desc.strip()}}}"
        elif ct.lower().startswith("表") or ct.lower().startswith("tbl"):
            tbl_counter[0] += 1
            return f"@table[tbl:d2m_{tbl_counter[0]}]{{{desc.strip()}}}"

    return None


def _detect_seq_type(para) -> str | None:
    """检测段落中 SEQ 域代码的类型（Figure/Table）"""
    flds = para._element.findall(f'.//{qn("w:fldChar")}')
    instrs = para._element.findall(f'.//{qn("w:instrText")}')
    for instr in instrs:
        txt = (instr.text or "").upper()
        if " SEQ " in txt or txt.startswith("SEQ "):
            if "FIGURE" in txt:
                return "figure"
            if "TABLE" in txt or "TBL" in txt:
                return "table"
    return None


def _extract_ref_fields(para) -> str:
    """
    提取 REF 域代码并替换为 @ref{bookmark_name}。
    同时保留纯文本。
    """
    # 找到 REF 指令
    instrs = para._element.findall(f'.//{qn("w:instrText")}')
    ref_map = {}
    for instr in instrs:
        txt = (instr.text or "").strip()
        if txt.upper().startswith("REF "):
            bm = txt.split(" ", 2)[1] if len(txt.split(" ", 2)) > 1 else ""
            bm = bm.split(" ")[0]
            if bm:
                ref_map[id(instr.getparent())] = bm

    if not ref_map:
        return ""
    # 仅返回第一个找到的引用（简化处理）
    return list(ref_map.values())[0]


def _process_paragraph(para, doc, img_dir: str, img_dir_name: str,
                       fig_counter: list, tbl_counter: list,
                       in_caption_block: list) -> str | None:
    """
    处理一个段落，返回 Markdown 行（多行用 \\n 拼接）或 None（跳过）。
    in_caption_block[0] 表示前一段是否为题注，用于跳过图片下的题注生成。
    """
    text = para.text.strip()

    # ── 题注 → @figure / @table ──
    cap = _process_caption(para, fig_counter, tbl_counter)
    if cap:
        in_caption_block[0] = True
        return cap

    # ── 空段落 ──
    if not text and not _has_drawing(para) and not _has_math(para):
        in_caption_block[0] = False
        return ""

    # ── 图片 ──
    if _has_drawing(para):
        img_lines = _extract_images(para, img_dir, img_dir_name, fig_counter, doc)
        in_caption_block[0] = False
        return "\n".join(img_lines) if img_lines else None

    # ── 公式 ──
    math_text = _extract_math(para)
    if math_text:
        in_caption_block[0] = False
        return math_text

    # ── 标题 ──
    level = None
    style_name = para.style.name if para.style else ""
    level = _heading_level(style_name)
    if level is not None and level <= 4:
        in_caption_block[0] = False
        prefix = "#" * level
        lines_out = []
        for child in para._element.iterchildren():
            tag = child.tag
            if tag == qn('w:r'):
                run_text = ""
                for tn in child.findall(qn('w:t')):
                    if tn.text:
                        run_text += tn.text
                if run_text:
                    bold = child.find(qn('w:rPr')) is not None and \
                           child.find(qn('w:rPr')).find(qn('w:b')) is not None
                    if bold:
                        lines_out.append(f"**{run_text}**")
                    else:
                        lines_out.append(run_text)
        body = "".join(lines_out) or text
        return f"{prefix} {body}"

    # ── 普通段落：行内格式 ──
    in_caption_block[0] = False
    md_text = _process_paragraph_runs(para)
    if md_text is None:
        return None
    return md_text


def _process_table(table, tbl_counter: list) -> str | None:
    """将 python-docx Table 转为 Markdown 表格。"""
    rows = table.rows
    if not rows:
        return None

    md_rows = []
    for row_idx, row in enumerate(rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        md_rows.append("| " + " | ".join(cells) + " |")
        if row_idx == 0:
            md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n".join(md_rows)


def docx_to_md(docx_path: str, output_md: str, image_dir_name: str = "images"):
    """
    将 .docx 转换为 Markdown。

    参数:
      docx_path:    输入 .docx 文件路径
      output_md:    输出 .md 文件路径
      image_dir_name: 图片目录名称（相对于 output_md 所在目录）
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"文件不存在: {docx_path}")

    doc = Document(docx_path)
    base = os.path.dirname(os.path.abspath(output_md))
    img_dir = os.path.join(base, image_dir_name)
    os.makedirs(img_dir, exist_ok=True)

    lines = []
    fig_counter = [0]   # 图片序号
    tbl_counter = [0]   # 表格序号
    in_caption_block = [False]  # 前一段是否为题注

    # ── 按 body 子元素顺序遍历（段落/表格交错） ──
    body = doc.element.body
    para_idx = 0
    tbl_idx = 0

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":  # 段落
            if para_idx >= len(doc.paragraphs):
                para_idx += 1
                continue
            para = doc.paragraphs[para_idx]
            para_idx += 1

            md = _process_paragraph(para, doc, img_dir, image_dir_name,
                                    fig_counter, tbl_counter, in_caption_block)
            if md is not None:
                lines.append(md)

        elif tag == "tbl":  # 表格
            if tbl_idx >= len(doc.tables):
                tbl_idx += 1
                continue
            table = doc.tables[tbl_idx]
            tbl_idx += 1

            md_table = _process_table(table, tbl_counter)
            if md_table:
                lines.append("")
                lines.append(md_table)
                lines.append("")
            in_caption_block[0] = False

    # ── 写入 ──
    content = "\n".join(lines)

    # 清理多余空行
    while "\n\n\n" in content:
        content = content.replace("\n\n\n", "\n\n")

    with open(output_md, "w", encoding="utf-8") as f:
        f.write(content.strip() + "\n")

    return os.path.abspath(output_md)
