"""公式处理：块公式编号 + 行内公式"""
from lxml import etree
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from .utils import W_NS, next_bookmark_id, add_omath_display, add_omath_inline
from .crossref import make_complex_field


def make_formula_block(doc, latex: str, chapter: int, eq_seq: int,
                       bookmark_name: str = "", styleref_level: int = 1):
    """创建块公式（3 列无边框表格：空 | 公式 | 编号）。

    Args:
        latex: LaTeX 公式文本
        chapter: 当前章号
        eq_seq:  公式序号
        bookmark_name: 书签名（空则不生成书签）
        styleref_level: 章号引用层级
    """
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci in range(3):
        c = tbl.cell(0, ci)
        c.width = Cm(3.0 if ci != 1 else 9.0)
        tcPr = c._tc.get_or_add_tcPr()
        for o in tcPr.findall(qn("w:tcBorders")):
            tcPr.remove(o)
        tcPr.append(parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0"/>'
            f'<w:left w:val="none" w:sz="0"/>'
            f'<w:bottom w:val="none" w:sz="0"/>'
            f'<w:right w:val="none" w:sz="0"/>'
            f'</w:tcBorders>'))

    # 中格：公式
    cm = tbl.cell(0, 1)
    cm.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_omath_display(cm.paragraphs[0], latex)

    # 右格：编号 (1-1)
    cr = tbl.cell(0, 2)
    cp = cr.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if bookmark_name:
        bm_id = next_bookmark_id()
        bm_s = etree.SubElement(cp._element, f'{{{W_NS}}}bookmarkStart')
        bm_s.set(f'{{{W_NS}}}id', str(bm_id))
        bm_s.set(f'{{{W_NS}}}name', bookmark_name)

    etree.SubElement(etree.SubElement(cp._element, f'{{{W_NS}}}r'),
                     f'{{{W_NS}}}t').text = '('
    make_complex_field(cp, [
        ('fldChar', 'begin'),
        ('instr', f' STYLEREF {styleref_level} \\s '),
        ('fldChar', 'separate'), ('text', str(chapter)), ('fldChar', 'end'),
    ])
    etree.SubElement(etree.SubElement(cp._element, f'{{{W_NS}}}r'),
                     f'{{{W_NS}}}t').text = '\u2013'
    make_complex_field(cp, [
        ('fldChar', 'begin'),
        ('instr', ' SEQ ( \\* ARABIC \\s 1 '),
        ('fldChar', 'separate'), ('text', str(eq_seq)), ('fldChar', 'end'),
    ])
    etree.SubElement(etree.SubElement(cp._element, f'{{{W_NS}}}r'),
                     f'{{{W_NS}}}t').text = ')'

    if bookmark_name:
        bm_e = etree.SubElement(cp._element, f'{{{W_NS}}}bookmarkEnd')
        bm_e.set(f'{{{W_NS}}}id', str(bm_id))

    return tbl
