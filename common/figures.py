"""图片处理：图片嵌入 + 题注"""
import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .utils import set_para_style
from .crossref import add_caption


def make_image_block(doc, src: str, width_pt: int = 300,
                     image_style_id: str = "",
                     caption_label="图", bookmark_name="",
                     caption_text="", caption_style_name="Caption",
                     styleref_level=1):
    """创建图片段落（居中）+ 下方题注。

    Args:
        doc: python-docx Document
        src: 图片文件路径
        image_style_id: 图片段落样式 ID
        caption_label: "图" or "Figure"
        bookmark_name: 题注书签（用于交叉引用）
        caption_text: 题注描述文字
        caption_style_name: 题注样式名
        styleref_level: 章号引用层级
    """
    # 图片段落
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_style_id:
        set_para_style(p, image_style_id)

    if src and os.path.exists(src):
        try:
            run = p.add_run("")
            run.add_picture(src, width=Pt(width_pt))
        except Exception as e:
            print(f"[Image error: {e}]")
            p.add_run(f"[图片: {src}]")
    else:
        p.add_run("[图 占位]")

    # 题注
    if bookmark_name:
        add_caption(doc, caption_label, bookmark_name, caption_text,
                    caption_style_name=caption_style_name,
                    styleref_level=styleref_level)

    return p
