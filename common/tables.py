"""表格处理：三线表 + 单元格样式 + 题注 + 行内公式"""
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from .utils import set_para_style, apply_font, set_cell_border, add_omath_inline
from .crossref import add_caption

RE_INLINE = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|\$[^$]+\$)')


def _fill_cell_paragraph(cp, text: str, font_cn="宋体", font_size=10.5, bold=False):
    """填充单元格段落，支持 **粗体** *斜体* $行内公式$。"""
    for seg in RE_INLINE.split(text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            r = cp.add_run(seg[2:-2])
            apply_font(r, bold=True, cn="黑体", sz=font_size)
        elif seg.startswith("*") and seg.endswith("*"):
            r = cp.add_run(seg[1:-1])
            apply_font(r, italic=True, cn="楷体", sz=font_size)
        elif seg.startswith("$") and seg.endswith("$"):
            add_omath_inline(cp, seg[1:-1])
        else:
            r = cp.add_run(seg)
            apply_font(r, cn=font_cn, sz=font_size, bold=bold)


def make_three_line_table(doc, rows: list, cell_style_id: str,
                          header_font_cn="黑体", body_font_cn="宋体",
                          font_size=10.5,
                          border_top_sz=None, border_mid_sz=None, border_bot_sz=None,
                          caption_label="表", bookmark_name="",
                          caption_text="", caption_style_name="Caption",
                          styleref_level=1):
    """创建三线表 + 上方题注。

    Args:
        cell_style_id: 单元格段落样式 ID
        header_font_cn: 表头中文字体
        body_font_cn:  表身中文字体
        caption_label: "表" or "Table"
        bookmark_name: 题注书签
        caption_text:  题注描述文字
        caption_style_name: 题注样式名
        styleref_level: 章号引用层级
    """
    # 题注（表上方）
    if bookmark_name:
        add_caption(doc, caption_label, bookmark_name, caption_text,
                    caption_style_name=caption_style_name,
                    styleref_level=styleref_level)

    n = max(len(r) for r in rows) if rows else 1
    tbl = doc.add_table(rows=len(rows), cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sz_t = border_top_sz or int(1.5 * 8)
    sz_m = border_mid_sz or int(0.75 * 8)
    sz_b = border_bot_sz or int(1.5 * 8)

    for ri, rd in enumerate(rows):
        for ci in range(n):
            c = tbl.cell(ri, ci)
            c.text = ""
            cp = c.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_style(cp, cell_style_id)

            cell_text = rd[ci] if ci < len(rd) else ""
            is_header = (ri == 0)
            _fill_cell_paragraph(cp, cell_text,
                                 font_cn=header_font_cn if is_header else body_font_cn,
                                 font_size=font_size, bold=is_header)

            if is_header:
                set_cell_border(c, top={"sz": sz_t}, bottom={"sz": sz_m})
            elif ri == len(rows) - 1:
                set_cell_border(c, bottom={"sz": sz_b})
            else:
                set_cell_border(c)
    return tbl
